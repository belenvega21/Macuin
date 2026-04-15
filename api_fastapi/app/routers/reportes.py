from fastapi import APIRouter, Depends, Query
from typing import Optional
from sqlalchemy.orm import Session
from app.data.database import SessionLocal
from app.data.models.autoparte import Autoparte
from app.data.models.pedido import Pedido
from app.data.models.detalle_pedido import DetallePedido
from app.data.models.usuario import Usuario
import pandas as pd
from datetime import datetime
from fastapi.responses import FileResponse
from docx import Document 
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from collections import defaultdict

router = APIRouter(prefix="/reportes", tags=["Reportes"])


def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


# inventario json
@router.get("/inventario")
def reporte_inventario(db: Session = Depends(get_db)):
    return db.query(Autoparte).all()


# total de productos
@router.get("/total-productos")
def total_productos(db: Session = Depends(get_db)):
    total = db.query(Autoparte).count()
    return {"total_productos": total}


# pedidos
@router.get("/pedidos")
def total_pedidos(db: Session = Depends(get_db)):
    total = db.query(Pedido).count()
    return {"total_pedidos": total}


# mas vendidos
@router.get("/mas-vendidos")
def productos_mas_vendidos(db: Session = Depends(get_db)):
    detalles = db.query(DetallePedido).all()

    conteo = {}
    for d in detalles:
        conteo[d.autoparte_id] = conteo.get(d.autoparte_id, 0) + d.cantidad

    return conteo


# clientes json
@router.get("/clientes")
def reporte_clientes(db: Session = Depends(get_db)):
    usuarios = db.query(Usuario).all()
    return [
        {
            "id": u.id,
            "nombre": u.nombre,
            "email": u.email,
            "telefono": u.telefono,
            "rol": u.rol
        }
        for u in usuarios
    ]


# inventario excel
@router.get("/inventario/excel")
def exportar_inventario_excel(db: Session = Depends(get_db)):
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    productos = db.query(Autoparte).all()

    data = [
        {"SKU": p.sku or "N/A", "Nombre": p.nombre, "Marca": p.marca, "Categoría": p.categoria, "Precio": p.precio, "Stock": p.stock}
        for p in productos
    ]

    df = pd.DataFrame(data)
    import tempfile, os
    file_path = os.path.join(tempfile.gettempdir(), "reporte_inventario.xlsx")
    df.to_excel(file_path, index=False, startrow=2)

    from openpyxl import load_workbook
    wb = load_workbook(file_path)
    ws = wb.active
    ws.title = "Inventario"

    header_font = Font(name="Calibri", bold=True, size=14, color="333333")
    sub_font = Font(name="Calibri", size=10, color="666666")
    col_header_font = Font(name="Calibri", bold=True, size=10, color="FFFFFF")
    col_header_fill = PatternFill(start_color="333333", end_color="333333", fill_type="solid")
    even_fill = PatternFill(start_color="F5F5F5", end_color="F5F5F5", fill_type="solid")
    border = Border(bottom=Side(style="thin", color="E0E0E0"))
    currency_fmt = '"$"#,##0.00'

    ws.merge_cells("A1:F1")
    ws["A1"] = "MACUIN AUTOPARTES — Reporte de Inventario"
    ws["A1"].font = header_font
    ws["A2"] = f"Total de referencias: {len(productos)} piezas únicas"
    ws["A2"].font = sub_font

    for col_idx in range(1, 7):
        cell = ws.cell(row=3, column=col_idx)
        cell.font = col_header_font
        cell.fill = col_header_fill
        cell.alignment = Alignment(horizontal="center")

    for row_idx in range(4, ws.max_row + 1):
        if row_idx % 2 == 0:
            for col_idx in range(1, 7):
                ws.cell(row=row_idx, column=col_idx).fill = even_fill
        for col_idx in range(1, 7):
            ws.cell(row=row_idx, column=col_idx).border = border
        ws.cell(row=row_idx, column=5).number_format = currency_fmt

    ws.column_dimensions["A"].width = 16
    ws.column_dimensions["B"].width = 30
    ws.column_dimensions["C"].width = 16
    ws.column_dimensions["D"].width = 16
    ws.column_dimensions["E"].width = 14
    ws.column_dimensions["F"].width = 10

    wb.save(file_path)
    return FileResponse(file_path, filename="reporte_inventario.xlsx")


# inventario word
@router.get("/inventario/word")
def exportar_inventario_word(db: Session = Depends(get_db)):
    from docx.shared import Inches, Pt, RGBColor, Cm, Emu
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    productos = db.query(Autoparte).all()

    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # Title Block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run('MACUIN')
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)
    run = title.add_run('  AUTOPARTES')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xb0, 0xb0, 0xb0)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = subtitle.add_run('Catálogo de Inventario')
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run = subtitle.add_run(f'\nDocumento Oficial  •  Generado: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Summary
    total_valor = sum(p.precio * p.stock for p in productos)
    p_info = doc.add_paragraph()
    run = p_info.add_run(f'Total de referencias: ')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    run = p_info.add_run(f'{len(productos)} piezas únicas')
    run.font.size = Pt(10)
    run.bold = True
    run = p_info.add_run(f'   |   Valoración Total: ')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    run = p_info.add_run(f'${total_valor:,.2f}')
    run.font.size = Pt(10)
    run.bold = True
    run.font.color.rgb = RGBColor(0x10, 0xb9, 0x81)

    doc.add_paragraph('')

    # Table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['SKU', 'Nombre / Descripción', 'Marca', 'Categoría', 'Precio', 'Stock']
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = 'Calibri'
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {qn('w:fill'): '1e293b', qn('w:val'): 'clear'})
        shading.append(shading_elm)

    for idx, prod in enumerate(productos):
        row = table.add_row().cells
        row[0].text = prod.sku or 'N/A'
        row[1].text = prod.nombre
        row[2].text = prod.marca or ''
        row[3].text = prod.categoria or ''
        row[4].text = f'${prod.precio:,.2f}'
        stock_text = f'{prod.stock} pz'
        row[5].text = stock_text
        
        # Alternating row colors
        if idx % 2 == 0:
            for cell in row:
                shading = cell._element.get_or_add_tcPr()
                shading_elm = shading.makeelement(qn('w:shd'), {qn('w:fill'): 'f8fafc', qn('w:val'): 'clear'})
                shading.append(shading_elm)
        
        # Font styling for all cells
        for j, cell in enumerate(row):
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'
                    if j == 4:
                        run.bold = True
                    if j == 5:
                        run.font.color.rgb = RGBColor(0x10, 0xb9, 0x81) if prod.stock > 10 else RGBColor(0xef, 0x44, 0x44)
                        run.bold = True

    # Column widths
    widths = [Cm(2.5), Cm(6), Cm(3), Cm(3), Cm(2.5), Cm(2)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    import tempfile, os
    file_path = os.path.join(tempfile.gettempdir(), 'reporte_inventario.docx')
    doc.save(file_path)
    return FileResponse(file_path, filename='reporte_inventario.docx')


# inventario pdf
@router.get("/inventario/pdf")
def exportar_inventario_pdf(db: Session = Depends(get_db)):
    from reportlab.platypus import Image as RLImage
    from PIL import Image, ImageDraw
    import requests
    from io import BytesIO
    import os

    productos = db.query(Autoparte).all()
    import tempfile, os
    file_path = os.path.join(tempfile.gettempdir(), "reporte_inventario.pdf")
    
    # Modern Document Margins & Title
    doc = SimpleDocTemplate(file_path, rightMargin=30, leftMargin=30, topMargin=50, bottomMargin=50, title="Directorio de Inventario", author="MACUIN Autopartes")
    styles = getSampleStyleSheet()

    elementos = []
    
    # Encabezado (Modern Minimalist)
    header_data = [
        [Paragraph("<font size=24><b>MACUIN</b></font><br/><font color='#b0b0b0' size=10>AUTOPARTES</font>", styles["Normal"]), 
         Paragraph("<font size=14><b>Catálogo de Inventario</b></font><br/><font color='#888888'>Documento Oficial</font>", styles["Normal"])]
    ]
    header_table = Table(header_data, colWidths=[300, 230])
    header_table.setStyle(TableStyle([
        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('BOTTOMPADDING', (0,0), (-1,-1), 15),
    ]))
    elementos.append(header_table)
    elementos.append(Spacer(1, 10))

    elementos.append(Paragraph(f"<font color='#3b82f6'><b>Total de referencias registradas:</b></font> {len(productos)} piezas únicas", styles["Normal"]))
    elementos.append(Spacer(1, 20))

    # Definir Helper de Imagen con Bordes Redondeados
    def fetch_rounded_image(url_or_path, size=40, radius=10):
        try:
            if not url_or_path:
                raise Exception("No image")
            
            if url_or_path.startswith("http"):
                # Handle localhost fallback if needed
                fetch_url = url_or_path.replace("localhost:8001", "127.0.0.1:8000")
            elif url_or_path.startswith("/"):
                fetch_url = f"http://127.0.0.1:8000{url_or_path}"
            else:
                fetch_url = f"http://127.0.0.1:8000/{url_or_path}"
                
            resp = requests.get(fetch_url, timeout=2)
            img_data = BytesIO(resp.content)
            img = Image.open(img_data).convert("RGBA")
            
            # Resize
            img = img.resize((size, size), Image.Resampling.LANCZOS)
            
            # Mask
            mask = Image.new("L", (size, size), 0)
            draw = ImageDraw.Draw(mask)
            draw.rounded_rectangle((0, 0, size, size), radius=radius, fill=255)
            
            # Apply mask & white background
            bg = Image.new("RGBA", (size, size), (255, 255, 255, 255))
            bg.paste(img, (0,0), mask=mask)
            
            img_io = BytesIO()
            bg.save(img_io, format="PNG")
            img_io.seek(0)
            
            rl_img = RLImage(img_io, width=size, height=size, mask='auto')
            return rl_img
        except Exception:
            return Paragraph("<font color='#cccccc'><i>S/I</i></font>", styles["Normal"])

    # Tabla Inventario Moderno
    tabla_data = [["Visual", "Detalle de Producto", "Marca", "Stock", "Precio Público"]]
    
    for p in productos:
        rl_img = fetch_rounded_image(p.imagen, size=40, radius=8)
        
        detalle = Paragraph(f"<b>{p.nombre}</b><br/><font size=8 color='#888888'>SKU: {p.sku or 'N/A'}</font>", styles["Normal"])
        marca = Paragraph(f"<font size=9>{p.marca}</font>", styles["Normal"])
        stock_txt = f"<font color='#10b981'><b>{p.stock} pz</b></font>" if p.stock > 10 else f"<font color='#ef4444'><b>{p.stock} pz</b></font>"
        stock = Paragraph(stock_txt, styles["Normal"])
        precio = Paragraph(f"<b>${p.precio:,.2f}</b>", styles["Normal"])
        
        tabla_data.append([rl_img, detalle, marca, stock, precio])

    t = Table(tabla_data, colWidths=[60, 240, 80, 70, 80])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#f8fafc")),
        ('TEXTCOLOR', (0,0), (-1,0), colors.HexColor("#334155")),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0,0), (-1,0), 12),
        ('TOPPADDING', (0,0), (-1,0), 12),
        ('LINEBELOW', (0,0), (-1,0), 1, colors.HexColor("#e2e8f0")),
        ('ALIGN', (3,1), (-1,-1), 'RIGHT'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('LINEBELOW', (0,1), (-1,-1), 0.5, colors.HexColor("#f1f5f9")),
        ('PADDING', (0,1), (-1,-1), 8),
    ]))
    elementos.append(t)

    doc.build(elementos)

    return FileResponse(file_path, media_type="application/pdf", headers={"Content-Disposition": "inline; filename=\"reporte_inventario_pro.pdf\""})


# clientes excel
@router.get("/clientes/excel")
def exportar_clientes_excel(db: Session = Depends(get_db)):
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    usuarios = db.query(Usuario).all()
    data = [{"ID": u.id, "Nombre": u.nombre, "Email": u.email, "Teléfono": u.telefono or "N/A", "Rol": u.rol} for u in usuarios]
    df = pd.DataFrame(data)
    import tempfile, os
    file_path = os.path.join(tempfile.gettempdir(), "reporte_clientes.xlsx")
    df.to_excel(file_path, index=False, startrow=2)
    from openpyxl import load_workbook
    wb = load_workbook(file_path)
    ws = wb.active
    ws.title = "Clientes"
    ws.merge_cells("A1:E1")
    ws["A1"] = "MACUIN AUTOPARTES — Directorio de Clientes"
    ws["A1"].font = Font(name="Calibri", bold=True, size=14, color="333333")
    ws["A2"] = f"N° de socios: {len(usuarios)} registrados"
    ws["A2"].font = Font(name="Calibri", size=10, color="666666")
    hdr_font = Font(name="Calibri", bold=True, size=10, color="FFFFFF")
    hdr_fill = PatternFill(start_color="333333", end_color="333333", fill_type="solid")
    even_fill = PatternFill(start_color="F5F5F5", end_color="F5F5F5", fill_type="solid")
    border = Border(bottom=Side(style="thin", color="E0E0E0"))
    for c in range(1, 6):
        cell = ws.cell(row=3, column=c)
        cell.font = hdr_font
        cell.fill = hdr_fill
        cell.alignment = Alignment(horizontal="center")
    for r in range(4, ws.max_row + 1):
        if r % 2 == 0:
            for c in range(1, 6): ws.cell(row=r, column=c).fill = even_fill
        for c in range(1, 6): ws.cell(row=r, column=c).border = border
    ws.column_dimensions["A"].width = 8
    ws.column_dimensions["B"].width = 25
    ws.column_dimensions["C"].width = 30
    ws.column_dimensions["D"].width = 16
    ws.column_dimensions["E"].width = 12
    wb.save(file_path)
    return FileResponse(file_path, filename="reporte_clientes.xlsx")


# clientes word
@router.get("/clientes/word")
def exportar_clientes_word(db: Session = Depends(get_db)):
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    usuarios = db.query(Usuario).all()
    
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # Title
    title = doc.add_paragraph()
    run = title.add_run('MACUIN')
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)
    run = title.add_run('  AUTOPARTES')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xb0, 0xb0, 0xb0)

    subtitle = doc.add_paragraph()
    run = subtitle.add_run('Directorio de Clientes')
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run = subtitle.add_run(f'\nDocumento Oficial  •  Generado: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    admins = len([u for u in usuarios if u.rol == 'admin'])
    clients = len([u for u in usuarios if u.rol != 'admin'])
    p_info = doc.add_paragraph()
    run = p_info.add_run(f'N° de socios: {len(usuarios)} registrados  |  Clientes: {clients}  |  Administradores: {admins}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    doc.add_paragraph('')

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['ID', 'Nombre / Razón Social', 'Contacto', 'Teléfono', 'Rol']
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = 'Calibri'
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {qn('w:fill'): '1e293b', qn('w:val'): 'clear'})
        shading.append(shading_elm)

    for idx, u in enumerate(usuarios):
        row = table.add_row().cells
        row[0].text = f'#{u.id}'
        row[1].text = u.nombre
        row[2].text = u.email
        row[3].text = u.telefono or 'N/A'
        row[4].text = u.rol.upper()
        
        if idx % 2 == 0:
            for cell in row:
                shading = cell._element.get_or_add_tcPr()
                shading_elm = shading.makeelement(qn('w:shd'), {qn('w:fill'): 'f8fafc', qn('w:val'): 'clear'})
                shading.append(shading_elm)
        
        for j, cell in enumerate(row):
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'
                    if j == 4:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0x3b, 0x82, 0xf6) if u.rol == 'admin' else RGBColor(0x64, 0x74, 0x8b)

    widths = [Cm(1.5), Cm(5), Cm(5.5), Cm(3), Cm(3)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    import tempfile, os
    file_path = os.path.join(tempfile.gettempdir(), 'reporte_clientes.docx')
    doc.save(file_path)
    return FileResponse(file_path, filename='reporte_clientes.docx')


# clientes pdf
@router.get("/clientes/pdf")
def exportar_clientes_pdf(db: Session = Depends(get_db)):
    usuarios = db.query(Usuario).all()

    import tempfile, os
    file_path = os.path.join(tempfile.gettempdir(), "reporte_clientes.pdf")
    doc = SimpleDocTemplate(file_path, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40, title="Reporte de Clientes", author="MACUIN Autopartes")
    styles = getSampleStyleSheet()

    elementos = []
    
    header_data = [
        [Paragraph("<b>MACUIN AUTOPARTES</b>", styles["Heading1"]), Paragraph("<font size=12><b>Directorio de Clientes</b></font>", styles["Normal"])]
    ]
    header_table = Table(header_data, colWidths=[300, 200])
    header_table.setStyle(TableStyle([
        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
        ('BACKGROUND', (1, 0), (1, 0), colors.HexColor("#f0f0f0")),
        ('PADDING', (1, 0), (1, 0), 10),
    ]))
    elementos.append(header_table)
    elementos.append(Spacer(1, 20))

    elementos.append(Paragraph(f"<b>N° de socios:</b> {len(usuarios)} registrados", styles["Normal"]))
    elementos.append(Spacer(1, 10))

    tabla_data = [["ID", "Nombre / Razón Social", "Contacto", "Teléfono", "Rol"]]
    for u in usuarios:
        tabla_data.append([
            f"#{u.id}",
            u.nombre,
            u.email,
            u.telefono or "N/A",
            u.rol
        ])

    t = Table(tabla_data, colWidths=[40, 160, 160, 80, 60])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#333333")),
        ('TEXTCOLOR', (0,0), (-1,0), colors.white),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('LINEBELOW', (0,0), (-1,-1), 0.5, colors.HexColor("#e0e0e0")),
        ('PADDING', (0,0), (-1,-1), 8),
    ]))
    elementos.append(t)

    doc.build(elementos)

    return FileResponse(file_path, media_type="application/pdf", headers={"Content-Disposition": "inline; filename=\"reporte_clientes.pdf\""})


# pedidos excel
@router.get("/pedidos/excel")
def exportar_pedidos_excel(start_date: Optional[str] = Query(None), end_date: Optional[str] = Query(None), db: Session = Depends(get_db)):
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from datetime import datetime
    query = db.query(Pedido)
    if start_date:
        query = query.filter(Pedido.fecha >= datetime.fromisoformat(start_date))
    if end_date:
        query = query.filter(Pedido.fecha <= datetime.fromisoformat(end_date))
    pedidos_db = query.all()
    data = []
    for p in pedidos_db:
        detalles = db.query(DetallePedido).filter(DetallePedido.pedido_id == p.id).all()
        data.append({"Pedido": f"ORD-{str(p.id).zfill(4)}", "Usuario ID": p.usuario_id, "Estado": p.estado, "Artículos": len(detalles)})
    df = pd.DataFrame(data)
    import tempfile, os
    file_path = os.path.join(tempfile.gettempdir(), "reporte_pedidos.xlsx")
    df.to_excel(file_path, index=False, startrow=2)
    from openpyxl import load_workbook
    wb = load_workbook(file_path)
    ws = wb.active
    ws.title = "Pedidos"
    ws.merge_cells("A1:D1")
    ws["A1"] = "MACUIN AUTOPARTES — Reporte Histórico de Pedidos"
    ws["A1"].font = Font(name="Calibri", bold=True, size=14, color="333333")
    ws["A2"] = f"Total: {len(pedidos_db)} pedidos procesados"
    ws["A2"].font = Font(name="Calibri", size=10, color="666666")
    hdr_font = Font(name="Calibri", bold=True, size=10, color="FFFFFF")
    hdr_fill = PatternFill(start_color="333333", end_color="333333", fill_type="solid")
    even_fill = PatternFill(start_color="F5F5F5", end_color="F5F5F5", fill_type="solid")
    border = Border(bottom=Side(style="thin", color="E0E0E0"))
    for c in range(1, 5):
        cell = ws.cell(row=3, column=c)
        cell.font = hdr_font
        cell.fill = hdr_fill
        cell.alignment = Alignment(horizontal="center")
    for r in range(4, ws.max_row + 1):
        if r % 2 == 0:
            for c in range(1, 5): ws.cell(row=r, column=c).fill = even_fill
        for c in range(1, 5): ws.cell(row=r, column=c).border = border
    ws.column_dimensions["A"].width = 16
    ws.column_dimensions["B"].width = 14
    ws.column_dimensions["C"].width = 20
    ws.column_dimensions["D"].width = 14
    wb.save(file_path)
    return FileResponse(file_path, filename="reporte_pedidos.xlsx")


# pedidos word
@router.get("/pedidos/word")
def exportar_pedidos_word(start_date: Optional[str] = Query(None), end_date: Optional[str] = Query(None), db: Session = Depends(get_db)):
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    query = db.query(Pedido)
    if start_date:
        query = query.filter(Pedido.fecha >= datetime.fromisoformat(start_date))
    if end_date:
        query = query.filter(Pedido.fecha <= datetime.fromisoformat(end_date))
    pedidos_db = query.all()
    usuarios = {u.id: u for u in db.query(Usuario).all()}
    autopartes_map = {a.id: a for a in db.query(Autoparte).all()}

    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # Title
    title = doc.add_paragraph()
    run = title.add_run('MACUIN')
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)
    run = title.add_run('  AUTOPARTES')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xb0, 0xb0, 0xb0)

    subtitle = doc.add_paragraph()
    run = subtitle.add_run('Reporte Histórico de Pedidos')
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    fecha_info = f'\nDocumento Oficial  •  Generado: {datetime.now().strftime("%d/%m/%Y %H:%M")}'
    if start_date or end_date:
        fecha_info += f'  •  Periodo: {(start_date or "Inicio").split("T")[0]} al {(end_date or "Actualidad").split("T")[0]}'
    run = subtitle.add_run(fecha_info)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Compute totals
    total_ingresos = 0
    for p in pedidos_db:
        if p.estado != 'cancelado':
            detalles = db.query(DetallePedido).filter(DetallePedido.pedido_id == p.id).all()
            for d in detalles:
                ap = autopartes_map.get(d.autoparte_id)
                if ap:
                    total_ingresos += ap.precio * d.cantidad

    p_info = doc.add_paragraph()
    run = p_info.add_run(f'Total: {len(pedidos_db)} pedidos procesados  |  Ingresos: ')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    run = p_info.add_run(f'${total_ingresos:,.2f}')
    run.font.size = Pt(10)
    run.bold = True
    run.font.color.rgb = RGBColor(0x10, 0xb9, 0x81)

    doc.add_paragraph('')

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Pedido', 'Cliente', 'Fecha', 'Estado', 'Artículos', 'Total']
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = 'Calibri'
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {qn('w:fill'): '1e293b', qn('w:val'): 'clear'})
        shading.append(shading_elm)

    for idx, p in enumerate(pedidos_db):
        detalles = db.query(DetallePedido).filter(DetallePedido.pedido_id == p.id).all()
        subtotal = sum((autopartes_map.get(d.autoparte_id).precio * d.cantidad) if autopartes_map.get(d.autoparte_id) else 0 for d in detalles)
        usr = usuarios.get(p.usuario_id)
        row = table.add_row().cells
        row[0].text = f'ORD-{str(p.id).zfill(4)}'
        row[1].text = usr.nombre if usr else f'ID:{p.usuario_id}'
        row[2].text = str(p.fecha).split('T')[0].split(' ')[0] if p.fecha else 'N/A'
        row[3].text = p.estado.upper().replace('_', ' ')
        row[4].text = f'{len(detalles)} items'
        row[5].text = f'${subtotal:,.2f}'
        
        # Status color
        estado_colors = {'RECIBIDO': 'dbeafe', 'EN PROCESO': 'fef3c7', 'EN RUTA': 'e0f2fe', 'ENTREGADO': 'd1fae5', 'CANCELADO': 'fee2e2'}
        status_bg = estado_colors.get(p.estado.upper().replace('_', ' '), 'ffffff')
        shading = row[3]._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {qn('w:fill'): status_bg, qn('w:val'): 'clear'})
        shading.append(shading_elm)
        
        if idx % 2 == 0:
            for ci in [0,1,2,4,5]:
                shading = row[ci]._element.get_or_add_tcPr()
                shading_elm = shading.makeelement(qn('w:shd'), {qn('w:fill'): 'f8fafc', qn('w:val'): 'clear'})
                shading.append(shading_elm)
        
        for j, cell in enumerate(row):
            for par in cell.paragraphs:
                for run in par.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'
                    if j == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
                    if j == 5:
                        run.bold = True

    widths = [Cm(2.5), Cm(4), Cm(2.5), Cm(3), Cm(2), Cm(3)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    # Total row at bottom
    doc.add_paragraph('')
    p_total = doc.add_paragraph()
    run = p_total.add_run(f'Ingresos Totales: ')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    run = p_total.add_run(f'${total_ingresos:,.2f}')
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = RGBColor(0x10, 0xb9, 0x81)

    import tempfile, os
    file_path = os.path.join(tempfile.gettempdir(), 'reporte_pedidos.docx')
    doc.save(file_path)
    return FileResponse(file_path, filename='reporte_pedidos.docx')


# pedidos pdf
@router.get("/pedidos/pdf")
def exportar_pedidos_pdf(start_date: Optional[str] = Query(None), end_date: Optional[str] = Query(None), db: Session = Depends(get_db)):
    from datetime import datetime
    query = db.query(Pedido)
    if start_date:
        query = query.filter(Pedido.fecha >= datetime.fromisoformat(start_date))
    if end_date:
        query = query.filter(Pedido.fecha <= datetime.fromisoformat(end_date))
    pedidos = query.all()

    import tempfile, os
    file_path = os.path.join(tempfile.gettempdir(), "reporte_pedidos.pdf")
    doc = SimpleDocTemplate(file_path, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40, title="Reporte de Pedidos", author="MACUIN Autopartes")
    styles = getSampleStyleSheet()

    elementos = []
    
    header_data = [
        [Paragraph("<b>MACUIN AUTOPARTES</b>", styles["Heading1"]), Paragraph("<font size=12><b>Reporte Histórico de Pedidos</b></font>", styles["Normal"])]
    ]
    header_table = Table(header_data, colWidths=[200, 300])
    header_table.setStyle(TableStyle([
        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
        ('BACKGROUND', (1, 0), (1, 0), colors.HexColor("#f0f0f0")),
        ('PADDING', (1, 0), (1, 0), 10),
    ]))
    elementos.append(header_table)
    elementos.append(Spacer(1, 20))

    elementos.append(Paragraph(f"<b>Total:</b> {len(pedidos)} pedidos procesados", styles["Normal"]))
    elementos.append(Spacer(1, 10))

    tabla_data = [["Pedido", "Usuario ID", "Estado", "Artículos"]]
    for p in pedidos:
        tabla_data.append([
            f"ORD-{str(p.id).zfill(4)}",
            f"ID: {p.usuario_id}",
            p.estado,
            f"{len(p.items)} items"
        ])

    t = Table(tabla_data, colWidths=[80, 100, 160, 100])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#333333")),
        ('TEXTCOLOR', (0,0), (-1,0), colors.white),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('LINEBELOW', (0,0), (-1,-1), 0.5, colors.HexColor("#e0e0e0")),
        ('PADDING', (0,0), (-1,-1), 8),
        ('ALIGN', (3,1), (3,-1), 'CENTER'),
    ]))
    elementos.append(t)

    doc.build(elementos)

    return FileResponse(file_path, media_type="application/pdf", headers={"Content-Disposition": "inline; filename=\"reporte_pedidos.pdf\""})


# recibo de pedido individual pdf
@router.get("/pedidos/{pedido_id}/recibo/pdf")
def exportar_recibo_pdf(pedido_id: int, db: Session = Depends(get_db)):
    pedido = db.query(Pedido).filter(Pedido.id == pedido_id).first()
    if not pedido: return {"error": "Pedido no encontrado"}
    
    usuario = db.query(Usuario).filter(Usuario.id == pedido.usuario_id).first()
    detalles = db.query(DetallePedido).filter(DetallePedido.pedido_id == pedido_id).all()
    autopartes = {a.id: a for a in db.query(Autoparte).all()}
    
    import tempfile, os
    file_path = os.path.join(tempfile.gettempdir(), f"recibo_orden_{pedido_id}.pdf")
    doc = SimpleDocTemplate(file_path, rightMargin=40, leftMargin=40, topMargin=50, bottomMargin=50, title=f"Recibo Pedido #{pedido_id}", author="MACUIN")
    styles = getSampleStyleSheet()
    elementos = []

    # Encabezado
    header_data = [
        [Paragraph("<font size=20><b>MACUIN AUTOPARTES</b></font>", styles["Normal"]), Paragraph("<font size=12 color='#888888'><b>Recibo de Compra</b></font>", styles["Normal"])]
    ]
    header_table = Table(header_data, colWidths=[300, 200])
    header_table.setStyle(TableStyle([
        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
        ('BOTTOMPADDING', (0,0), (-1,-1), 15),
    ]))
    elementos.append(header_table)

    # Info Cliente y Ticket
    fecha_str = pedido.fecha.strftime("%Y-%m-%d %H:%M") if pedido.fecha else "N/A"
    info_data = [
        [Paragraph(f"<b>Acreedor:</b><br/>{usuario.nombre if usuario else 'N/A'}<br/>{usuario.email if usuario else ''}", styles["Normal"]), 
         Paragraph(f"<b>Ticket M-ORD-{pedido_id}</b><br/>Fecha: {fecha_str}<br/>Estado: {pedido.estado.upper()}", styles["Normal"])]
    ]
    info_table = Table(info_data, colWidths=[250, 250])
    info_table.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'TOP'), ('TOPPADDING', (0,0), (-1,-1), 10)]))
    elementos.append(info_table)
    elementos.append(Spacer(1, 20))

    # Productos
    total = 0
    tabla_data = [["Producto / Descripción", "Cant.", "Precio U.", "Importe"]]
    for d in detalles:
        ap = autopartes.get(d.autoparte_id)
        if ap:
            sub = ap.precio * d.cantidad
            total += sub
            tabla_data.append([
                Paragraph(f"<b>{ap.nombre}</b><br/><font size=8>{ap.marca}</font>", styles["Normal"]),
                f"x{d.cantidad}",
                f"${ap.precio:,.2f}",
                f"${sub:,.2f}"
            ])
            
    t = Table(tabla_data, colWidths=[260, 60, 80, 100])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#1e293b")),
        ('TEXTCOLOR', (0,0), (-1,0), colors.white),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('PADDING', (0,0), (-1,-1), 10),
        ('LINEBELOW', (0,1), (-1,-1), 0.5, colors.HexColor("#e2e8f0")),
        ('ALIGN', (1,0), (-1,-1), 'RIGHT'),
        ('ALIGN', (0,0), (0,-1), 'LEFT'),
    ]))
    elementos.append(t)
    elementos.append(Spacer(1, 15))

    # Totales
    total_table = Table([
        ["Subtotal", f"${total:,.2f}"],
        ["Envío Vía", f"{pedido.paqueteria if pedido.paqueteria else 'Local'}"],
        [Paragraph("<b>TOTAL PAGADO</b>", styles["Normal"]), Paragraph(f"<b>${total:,.2f}</b>", styles["Normal"])]
    ], colWidths=[100, 100])
    total_table.setStyle(TableStyle([
        ('ALIGN', (0,0), (-1,-1), 'RIGHT'),
        ('LINEABOVE', (0,2), (1,2), 1, colors.black),
        ('TOPPADDING', (0,0), (-1,-1), 5),
    ]))
    
    layout = Table([["", total_table]], colWidths=[300, 200])
    elementos.append(layout)

    doc.build(elementos)
    return FileResponse(file_path, media_type="application/pdf", headers={"Content-Disposition": f"inline; filename=\"recibo_orden_{pedido_id}.pdf\""})


# ventas excel
@router.get("/ventas/excel")
def exportar_ventas_excel(start_date: Optional[str] = Query(None), end_date: Optional[str] = Query(None), db: Session = Depends(get_db)):
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    query = db.query(DetallePedido).join(Pedido, DetallePedido.pedido_id == Pedido.id)
    if start_date:
        query = query.filter(Pedido.fecha >= datetime.fromisoformat(start_date))
    if end_date:
        query = query.filter(Pedido.fecha <= datetime.fromisoformat(end_date))
    detalles = query.all()
    autopartes = {a.id: a for a in db.query(Autoparte).all()}
    data = []
    total_global = 0
    for d in detalles:
        ap = autopartes.get(d.autoparte_id)
        if ap:
            total = ap.precio * d.cantidad
            total_global += total
            data.append({"Autoparte": ap.nombre, "Precio Unit.": ap.precio, "Cant.": d.cantidad, "Subtotal": total})
    df = pd.DataFrame(data)
    import tempfile, os
    file_path = os.path.join(tempfile.gettempdir(), "reporte_ventas.xlsx")
    df.to_excel(file_path, index=False, startrow=2)
    from openpyxl import load_workbook
    wb = load_workbook(file_path)
    ws = wb.active
    ws.title = "Ventas"
    ws.merge_cells("A1:D1")
    ws["A1"] = "MACUIN AUTOPARTES — Reporte Analítico de Ventas"
    ws["A1"].font = Font(name="Calibri", bold=True, size=14, color="333333")
    ws["A2"] = f"Operaciones totales registradas: {len(detalles)} movimientos"
    ws["A2"].font = Font(name="Calibri", size=10, color="666666")
    hdr_font = Font(name="Calibri", bold=True, size=10, color="FFFFFF")
    hdr_fill = PatternFill(start_color="333333", end_color="333333", fill_type="solid")
    even_fill = PatternFill(start_color="F5F5F5", end_color="F5F5F5", fill_type="solid")
    border = Border(bottom=Side(style="thin", color="E0E0E0"))
    currency_fmt = '"$"#,##0.00'
    for c in range(1, 5):
        cell = ws.cell(row=3, column=c)
        cell.font = hdr_font
        cell.fill = hdr_fill
        cell.alignment = Alignment(horizontal="center")
    for r in range(4, ws.max_row + 1):
        if r % 2 == 0:
            for c in range(1, 5): ws.cell(row=r, column=c).fill = even_fill
        for c in range(1, 5): ws.cell(row=r, column=c).border = border
        ws.cell(row=r, column=2).number_format = currency_fmt
        ws.cell(row=r, column=4).number_format = currency_fmt
    total_row = ws.max_row + 2
    ws.cell(row=total_row, column=3).value = "TOTAL:"
    ws.cell(row=total_row, column=3).font = Font(name="Calibri", bold=True, size=12)
    ws.cell(row=total_row, column=4).value = total_global
    ws.cell(row=total_row, column=4).number_format = currency_fmt
    ws.cell(row=total_row, column=4).font = Font(name="Calibri", bold=True, size=12, color="333333")
    ws.column_dimensions["A"].width = 30
    ws.column_dimensions["B"].width = 16
    ws.column_dimensions["C"].width = 10
    ws.column_dimensions["D"].width = 20
    wb.save(file_path)
    return FileResponse(file_path, filename="reporte_ventas.xlsx")


# ventas word
@router.get("/ventas/word")
def exportar_ventas_word(start_date: Optional[str] = Query(None), end_date: Optional[str] = Query(None), db: Session = Depends(get_db)):
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    query = db.query(DetallePedido).join(Pedido, DetallePedido.pedido_id == Pedido.id)
    if start_date:
        query = query.filter(Pedido.fecha >= datetime.fromisoformat(start_date))
    if end_date:
        query = query.filter(Pedido.fecha <= datetime.fromisoformat(end_date))
    detalles = query.all()
    autopartes = {a.id: a for a in db.query(Autoparte).all()}

    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # Title
    title = doc.add_paragraph()
    run = title.add_run('MACUIN')
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)
    run = title.add_run('  AUTOPARTES')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xb0, 0xb0, 0xb0)

    subtitle = doc.add_paragraph()
    run = subtitle.add_run('Reporte Analítico de Ventas')
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    fecha_info = f'\nDocumento Oficial  •  Generado: {datetime.now().strftime("%d/%m/%Y %H:%M")}'
    if start_date or end_date:
        fecha_info += f'  •  Periodo: {(start_date or "Inicio").split("T")[0]} al {(end_date or "Actualidad").split("T")[0]}'
    run = subtitle.add_run(fecha_info)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph('')

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Autoparte', 'Precio Unitario', 'Cantidad', 'Subtotal']
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = 'Calibri'
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {qn('w:fill'): '1e293b', qn('w:val'): 'clear'})
        shading.append(shading_elm)

    total_global = 0
    for idx, d in enumerate(detalles):
        ap = autopartes.get(d.autoparte_id)
        if ap:
            total = ap.precio * d.cantidad
            total_global += total
            row = table.add_row().cells
            row[0].text = ap.nombre
            row[1].text = f'${ap.precio:,.2f}'
            row[2].text = f'x{d.cantidad}'
            row[3].text = f'${total:,.2f}'
            
            if idx % 2 == 0:
                for cell in row:
                    shading = cell._element.get_or_add_tcPr()
                    shading_elm = shading.makeelement(qn('w:shd'), {qn('w:fill'): 'f8fafc', qn('w:val'): 'clear'})
                    shading.append(shading_elm)
            
            for j, cell in enumerate(row):
                for par in cell.paragraphs:
                    for run in par.runs:
                        run.font.size = Pt(9)
                        run.font.name = 'Calibri'
                        if j == 3:
                            run.bold = True

    widths = [Cm(7), Cm(3.5), Cm(2.5), Cm(4)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    # Total
    doc.add_paragraph('')
    p_info = doc.add_paragraph()
    run = p_info.add_run(f'Operaciones registradas: {len(detalles)} movimientos')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    p_total = doc.add_paragraph()
    run = p_total.add_run('Ingresos Totales Globales: ')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    run = p_total.add_run(f'${total_global:,.2f}')
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = RGBColor(0x10, 0xb9, 0x81)

    import tempfile, os
    file_path = os.path.join(tempfile.gettempdir(), 'reporte_ventas.docx')
    doc.save(file_path)
    return FileResponse(file_path, filename='reporte_ventas.docx')


# ventas pdf
@router.get("/ventas/pdf")
def exportar_ventas_pdf(start_date: Optional[str] = Query(None), end_date: Optional[str] = Query(None), db: Session = Depends(get_db)):
    query = db.query(DetallePedido).join(Pedido, DetallePedido.pedido_id == Pedido.id)
    if start_date:
        query = query.filter(Pedido.fecha >= datetime.fromisoformat(start_date))
    if end_date:
        query = query.filter(Pedido.fecha <= datetime.fromisoformat(end_date))
    detalles = query.all()
    autopartes = {a.id: a for a in db.query(Autoparte).all()}

    import tempfile, os
    file_path = os.path.join(tempfile.gettempdir(), "reporte_ventas.pdf")
    doc = SimpleDocTemplate(file_path, rightMargin=30, leftMargin=30, topMargin=50, bottomMargin=50, title="Reporte Analítico de Ventas", author="MACUIN Autopartes")
    styles = getSampleStyleSheet()

    elementos = []
    
    header_data = [
        [Paragraph("<font size=24><b>MACUIN</b></font><br/><font color='#b0b0b0' size=10>AUTOPARTES</font>", styles["Normal"]), 
         Paragraph("<font size=14><b>Reporte Analítico de Ventas</b></font><br/><font color='#888888'>Documento Oficial</font>", styles["Normal"])]
    ]
    header_table = Table(header_data, colWidths=[300, 230])
    header_table.setStyle(TableStyle([
        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('BOTTOMPADDING', (0,0), (-1,-1), 15),
    ]))
    elementos.append(header_table)
    elementos.append(Spacer(1, 20))

    elementos.append(Paragraph(f"<font color='#8b5cf6'><b>Operaciones totales registradas:</b></font> {len(detalles)} movimientos", styles["Normal"]))
    elementos.append(Spacer(1, 15))

    tabla_data = [["Autoparte", "Precio Unit.", "Cant.", "Subtotal por Producto"]]
    total_global = 0
    for d in detalles:
        ap = autopartes.get(d.autoparte_id)
        if ap:
            total = ap.precio * d.cantidad
            total_global += total
            
            detalle = Paragraph(f"<b>{ap.nombre}</b>", styles["Normal"])
            
            tabla_data.append([
                detalle,
                Paragraph(f"<font color='#64748b'>${ap.precio:,.2f}</font>", styles["Normal"]),
                f"x{d.cantidad}",
                Paragraph(f"<b>${total:,.2f}</b>", styles["Normal"])
            ])

    t = Table(tabla_data, colWidths=[240, 100, 60, 130])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#f8fafc")),
        ('TEXTCOLOR', (0,0), (-1,0), colors.HexColor("#334155")),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0,0), (-1,0), 12),
        ('TOPPADDING', (0,0), (-1,0), 12),
        ('LINEBELOW', (0,0), (-1,0), 1, colors.HexColor("#e2e8f0")),
        ('ALIGN', (1,1), (-1,-1), 'RIGHT'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('LINEBELOW', (0,1), (-1,-1), 0.5, colors.HexColor("#f1f5f9")),
        ('PADDING', (0,1), (-1,-1), 10),
    ]))
    elementos.append(t)
    elementos.append(Spacer(1, 25))
    
    elementos.append(Paragraph(f"<font size=14>Ingresos Totales Globales:</font> <font size=18 color='#10b981'><b>${total_global:,.2f}</b></font>", styles["Normal"]))

    doc.build(elementos)

    return FileResponse(file_path, media_type="application/pdf", headers={"Content-Disposition": "inline; filename=\"reporte_ventas_pro.pdf\""})