from fastapi import APIRouter, Depends, Query
from sqlalchemy.orm import Session
from app.data.database import SessionLocal
from app.data.models.pedido import Pedido
from app.data.models.detalle_pedido import DetallePedido
from app.data.models.usuario import Usuario
from app.data.models.autoparte import Autoparte
import pandas as pd
from fastapi.responses import FileResponse
from docx import Document 
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from typing import Optional
from datetime import datetime

router = APIRouter(prefix="/reportes", tags=["Reportes Extra"])

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


# ================== CLIENTES ==================

@router.get("/clientes/excel")
def exportar_clientes_excel(db: Session = Depends(get_db)):
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    usuarios = db.query(Usuario).all()
    data = [{"ID": u.id, "Nombre": u.nombre, "Email": u.email, "Teléfono": u.telefono, "Rol": u.rol} for u in usuarios]
    df = pd.DataFrame(data)
    file_path = "/tmp/reporte_clientes.xlsx"
    df.to_excel(file_path, index=False, startrow=2)

    from openpyxl import load_workbook
    wb = load_workbook(file_path)
    ws = wb.active
    ws.title = "Directorio de Clientes"

    header_font = Font(name="Calibri", bold=True, size=14, color="333333")
    sub_font = Font(name="Calibri", size=10, color="666666")
    col_header_font = Font(name="Calibri", bold=True, size=10, color="FFFFFF")
    col_header_fill = PatternFill(start_color="333333", end_color="333333", fill_type="solid")
    even_fill = PatternFill(start_color="F5F5F5", end_color="F5F5F5", fill_type="solid")
    border = Border(bottom=Side(style="thin", color="E0E0E0"))

    ws.merge_cells("A1:E1")
    ws["A1"] = "MACUIN AUTOPARTES — Directorio de Cuentas"
    ws["A1"].font = header_font
    ws["A2"] = f"Total de cuentas registradas: {len(usuarios)}"
    ws["A2"].font = sub_font

    for col_idx in range(1, 6):
        cell = ws.cell(row=3, column=col_idx)
        cell.font = col_header_font
        cell.fill = col_header_fill
        cell.alignment = Alignment(horizontal="center")

    for row_idx in range(4, ws.max_row + 1):
        if row_idx % 2 == 0:
            for col_idx in range(1, 6):
                ws.cell(row=row_idx, column=col_idx).fill = even_fill
        for col_idx in range(1, 6):
            ws.cell(row=row_idx, column=col_idx).border = border

    ws.column_dimensions["A"].width = 10
    ws.column_dimensions["B"].width = 30
    ws.column_dimensions["C"].width = 30
    ws.column_dimensions["D"].width = 15
    ws.column_dimensions["E"].width = 15
    wb.save(file_path)
    return FileResponse(file_path, filename="reporte_clientes.xlsx")


@router.get("/clientes/word")
def exportar_clientes_word(db: Session = Depends(get_db)):
    usuarios = db.query(Usuario).all()
    doc = Document()
    doc.add_heading("MACUIN AUTOPARTES", 0)
    doc.add_heading("Directorio de Cuentas", level=2)
    doc.add_paragraph(f"Cuentas registradas: {len(usuarios)}")
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["ID", "Nombre", "Email", "Teléfono", "Rol"]
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        from docx.oxml.ns import qn
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {qn('w:fill'): '333333', qn('w:val'): 'clear'})
        shading.append(shading_elm)

    for u in usuarios:
        row = table.add_row().cells
        row[0].text = str(u.id)
        row[1].text = u.nombre
        row[2].text = u.email
        row[3].text = u.telefono or ""
        row[4].text = u.rol

    file_path = "/tmp/reporte_clientes.docx"
    doc.save(file_path)
    return FileResponse(file_path, filename="reporte_clientes.docx")


@router.get("/clientes/pdf")
def exportar_clientes_pdf(db: Session = Depends(get_db)):
    usuarios = db.query(Usuario).all()
    file_path = "/tmp/reporte_clientes.pdf"
    doc = SimpleDocTemplate(file_path, rightMargin=30, leftMargin=30, topMargin=50, bottomMargin=50)
    styles = getSampleStyleSheet()
    elementos = []
    
    header_data = [
        [Paragraph("<font size=24><b>MACUIN</b></font><br/><font color='#b0b0b0' size=10>AUTOPARTES</font>", styles["Normal"]), 
         Paragraph("<font size=14><b>Directorio de Cuentas</b></font><br/><font color='#888888'>Documento Oficial</font>", styles["Normal"])]
    ]
    header_table = Table(header_data, colWidths=[300, 230])
    header_table.setStyle(TableStyle([
        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('BOTTOMPADDING', (0,0), (-1,-1), 15),
    ]))
    elementos.append(header_table)
    elementos.append(Spacer(1, 20))
    elementos.append(Paragraph(f"<b>Cuentas registradas:</b> {len(usuarios)}", styles["Normal"]))
    elementos.append(Spacer(1, 15))

    tabla_data = [["ID", "Nombre", "Email", "Contacto", "Rol"]]
    for u in usuarios:
        tabla_data.append([str(u.id), u.nombre, u.email, u.telefono or "S/I", u.rol])

    t = Table(tabla_data, colWidths=[40, 160, 180, 80, 70])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#334155")),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0,0), (-1,-1), 8),
        ('TOPPADDING', (0,0), (-1,-1), 8),
        ('LINEBELOW', (0,0), (-1,-1), 0.5, colors.HexColor("#e2e8f0")),
    ]))
    elementos.append(t)
    doc.build(elementos)
    return FileResponse(file_path, media_type="application/pdf", headers={"Content-Disposition": "inline; filename=\"reporte_clientes.pdf\""})


# ================== PEDIDOS ==================

@router.get("/pedidos_historial/excel")
def exportar_pedidos_excel(start_date: Optional[str] = Query(None), end_date: Optional[str] = Query(None), db: Session = Depends(get_db)):
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    query = db.query(Pedido)
    if start_date:
        query = query.filter(Pedido.fecha >= datetime.fromisoformat(start_date))
    if end_date:
        query = query.filter(Pedido.fecha <= datetime.fromisoformat(end_date))
    pedidos = query.all()

    data = [{"ID Pedido": p.id, "ID Cliente": p.usuario_id, "Fecha": str(p.fecha), "Estado": p.estado, "Paquetería": p.paqueteria, "Rastreo": p.num_seguimiento} for p in pedidos]
    df = pd.DataFrame(data)
    file_path = "/tmp/reporte_pedidos.xlsx"
    df.to_excel(file_path, index=False, startrow=2)

    from openpyxl import load_workbook
    wb = load_workbook(file_path)
    ws = wb.active
    ws.title = "Historial Logístico"

    header_font = Font(name="Calibri", bold=True, size=14, color="333333")
    sub_font = Font(name="Calibri", size=10, color="666666")
    col_header_font = Font(name="Calibri", bold=True, size=10, color="FFFFFF")
    col_header_fill = PatternFill(start_color="333333", end_color="333333", fill_type="solid")
    even_fill = PatternFill(start_color="F5F5F5", end_color="F5F5F5", fill_type="solid")
    border = Border(bottom=Side(style="thin", color="E0E0E0"))

    ws.merge_cells("A1:F1")
    ws["A1"] = "MACUIN AUTOPARTES — Historial de Pedidos Logísticos"
    ws["A1"].font = header_font
    ws["A2"] = f"Total de operaciones: {len(pedidos)} (Filtro: {start_date or 'Inicio'} a {end_date or 'Fin'})"
    ws["A2"].font = sub_font

    for col_idx in range(1, 7):
        cell = ws.cell(row=3, column=col_idx)
        cell.font = col_header_font
        cell.fill = col_header_fill

    for row_idx in range(4, ws.max_row + 1):
        if row_idx % 2 == 0:
            for col_idx in range(1, 7):
                ws.cell(row=row_idx, column=col_idx).fill = even_fill
        for col_idx in range(1, 7):
            ws.cell(row=row_idx, column=col_idx).border = border

    ws.column_dimensions["A"].width = 12
    ws.column_dimensions["B"].width = 12
    ws.column_dimensions["C"].width = 25
    ws.column_dimensions["D"].width = 15
    ws.column_dimensions["E"].width = 20
    ws.column_dimensions["F"].width = 20
    wb.save(file_path)
    return FileResponse(file_path, filename="reporte_pedidos.xlsx")


@router.get("/pedidos_historial/word")
def exportar_pedidos_word(start_date: Optional[str] = Query(None), end_date: Optional[str] = Query(None), db: Session = Depends(get_db)):
    query = db.query(Pedido)
    if start_date:
        query = query.filter(Pedido.fecha >= datetime.fromisoformat(start_date))
    if end_date:
        query = query.filter(Pedido.fecha <= datetime.fromisoformat(end_date))
    pedidos = query.all()

    doc = Document()
    doc.add_heading("MACUIN AUTOPARTES", 0)
    doc.add_heading("Reporte de Pedidos", level=2)
    doc.add_paragraph(f"Operaciones: {len(pedidos)} | Fechas contempladas: {start_date or 'Inicio'} a {end_date or 'Fin'}")
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Ord.", "Usu.", "Fec.", "Est.", "Paq.", "Segum."]
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)

    for p in pedidos:
        row = table.add_row().cells
        row[0].text = str(p.id)
        row[1].text = str(p.usuario_id)
        row[2].text = str(p.fecha).split()[0]
        row[3].text = p.estado
        row[4].text = p.paqueteria or ""
        row[5].text = p.num_seguimiento or ""

    file_path = "/tmp/reporte_pedidos.docx"
    doc.save(file_path)
    return FileResponse(file_path, filename="reporte_pedidos.docx")


@router.get("/pedidos_historial/pdf")
def exportar_pedidos_pdf(start_date: Optional[str] = Query(None), end_date: Optional[str] = Query(None), db: Session = Depends(get_db)):
    query = db.query(Pedido)
    if start_date:
        query = query.filter(Pedido.fecha >= datetime.fromisoformat(start_date))
    if end_date:
        query = query.filter(Pedido.fecha <= datetime.fromisoformat(end_date))
    pedidos = query.all()

    file_path = "/tmp/reporte_pedidos.pdf"
    doc = SimpleDocTemplate(file_path, rightMargin=30, leftMargin=30, topMargin=50, bottomMargin=50)
    styles = getSampleStyleSheet()
    elementos = []
    
    header_data = [
        [Paragraph("<font size=24><b>MACUIN</b></font><br/><font color='#b0b0b0' size=10>AUTOPARTES</font>", styles["Normal"]), 
         Paragraph("<font size=14><b>Historial Logístico</b></font><br/><font color='#888888'>Documento Oficial</font>", styles["Normal"])]
    ]
    header_table = Table(header_data, colWidths=[300, 230])
    header_table.setStyle(TableStyle([('ALIGN', (1, 0), (1, 0), 'RIGHT'), ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'), ('BOTTOMPADDING', (0,0), (-1,-1), 15)]))
    elementos.append(header_table)
    elementos.append(Paragraph(f"<b>Operaciones:</b> {len(pedidos)} &nbsp;&nbsp; <b>Fechas:</b> {start_date or 'Toda la historia'} al {end_date or 'Actualidad'}", styles["Normal"]))
    elementos.append(Spacer(1, 15))

    tabla_data = [["Ticket", "Emisión", "Estatus", "Servicio", "Rastreo"]]
    for p in pedidos:
        tabla_data.append([str(p.id), str(p.fecha).split()[0], p.estado.upper(), p.paqueteria or "-", p.num_seguimiento or "-"])

    t = Table(tabla_data, colWidths=[60, 100, 100, 120, 150])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#0f172a")),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0,0), (-1,-1), 8),
        ('TOPPADDING', (0,0), (-1,-1), 8),
        ('LINEBELOW', (0,0), (-1,-1), 0.5, colors.HexColor("#e2e8f0")),
    ]))
    elementos.append(t)
    doc.build(elementos)
    return FileResponse(file_path, media_type="application/pdf", headers={"Content-Disposition": "inline; filename=\"reporte_pedidos.pdf\""})
